#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx2latex_converter 单元测试
"""

import sys
import tempfile
import unittest
from pathlib import Path
from docx import Document

# 添加 backend 目录到 Python 路径
backend_dir = Path(__file__).parent.parent
sys.path.insert(0, str(backend_dir))

from docx2latex_converter import Docx2LaTeXConverter


class TestExtractHeadings(unittest.TestCase):
    """测试 extract_content 方法 - 标题提取"""

    def test_extract_headings(self):
        """测试从 DOCX 提取标题"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            doc.add_heading('Introduction', level=1)
            doc.add_heading('Background', level=2)
            doc.add_heading('Related Work', level=3)
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)

            assert 'headings' in content
            assert len(content['headings']) == 3
            assert content['headings'][0]['text'] == 'Introduction'
            assert content['headings'][0]['level'] == 1
            assert content['headings'][1]['level'] == 2
            assert content['headings'][2]['level'] == 3
        finally:
            Path(doc_path).unlink(missing_ok=True)

    def test_extract_headings_only(self):
        """测试只包含标题的文档"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            doc.add_heading('Title', level=1)
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)

            assert len(content['headings']) == 1
            assert len(content['paragraphs']) == 0
        finally:
            Path(doc_path).unlink(missing_ok=True)


class TestExtractParagraphs(unittest.TestCase):
    """测试 extract_content 方法 - 段落提取"""

    def test_extract_paragraphs(self):
        """测试从 DOCX 提取段落"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            doc.add_paragraph('First paragraph text')
            doc.add_paragraph('Second paragraph text')
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)

            assert 'paragraphs' in content
            assert len(content['paragraphs']) == 2
            assert 'First paragraph text' in content['paragraphs']
            assert 'Second paragraph text' in content['paragraphs']
        finally:
            Path(doc_path).unlink(missing_ok=True)

    def test_extract_paragraphs_with_headings(self):
        """测试混合标题和段落"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            doc.add_heading('Section Title', level=1)
            doc.add_paragraph('Some paragraph content')
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)

            assert len(content['headings']) == 1
            assert len(content['paragraphs']) == 1
            assert content['paragraphs'][0] == 'Some paragraph content'
        finally:
            Path(doc_path).unlink(missing_ok=True)

    def test_empty_paragraphs_skipped(self):
        """测试空段落被跳过"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            doc.add_paragraph('Non-empty')
            doc.add_paragraph('')
            doc.add_paragraph('   ')
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)

            assert len(content['paragraphs']) == 1
            assert content['paragraphs'][0] == 'Non-empty'
        finally:
            Path(doc_path).unlink(missing_ok=True)


class TestExtractTables(unittest.TestCase):
    """测试 extract_content 方法 - 表格提取"""

    def test_extract_tables(self):
        """测试从 DOCX 提取表格"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            table = doc.add_table(rows=3, cols=2)
            table.cell(0, 0).text = 'A1'
            table.cell(0, 1).text = 'B1'
            table.cell(1, 0).text = 'A2'
            table.cell(1, 1).text = 'B2'
            table.cell(2, 0).text = 'A3'
            table.cell(2, 1).text = 'B3'
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)

            assert 'tables' in content
            assert len(content['tables']) == 1
            table_data = content['tables'][0]
            assert len(table_data) == 3
            assert table_data[0][0] == 'A1'
            assert table_data[0][1] == 'B1'
        finally:
            Path(doc_path).unlink(missing_ok=True)

    def test_extract_multiple_tables(self):
        """测试提取多个表格"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            doc.add_table(rows=2, cols=2)
            doc.add_table(rows=2, cols=3)
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)

            assert len(content['tables']) == 2
        finally:
            Path(doc_path).unlink(missing_ok=True)


class TestHeadingToLatex(unittest.TestCase):
    """测试标题到 LaTeX 的转换"""

    def test_section_heading(self):
        """测试一级标题转为 \\section"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            doc.add_heading('Chapter One', level=1)
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)
            latex = converter.convert_structure(content)

            assert r'\section{Chapter One}' in latex
        finally:
            Path(doc_path).unlink(missing_ok=True)

    def test_subsection_heading(self):
        """测试二级标题转为 \\subsection"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            doc.add_heading('Sub Section', level=2)
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)
            latex = converter.convert_structure(content)

            assert r'\subsection{Sub Section}' in latex
        finally:
            Path(doc_path).unlink(missing_ok=True)

    def test_subsubsection_heading(self):
        """测试三级标题转为 \\subsubsection"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            doc.add_heading('Sub Sub Section', level=3)
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)
            latex = converter.convert_structure(content)

            assert r'\subsubsection{Sub Sub Section}' in latex
        finally:
            Path(doc_path).unlink(missing_ok=True)


class TestParagraphToLatex(unittest.TestCase):
    """测试段落到 LaTeX 的转换"""

    def test_paragraph_to_latex(self):
        """测试段落转为 LaTeX"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            doc.add_paragraph('This is a paragraph.')
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)
            latex = converter.convert_structure(content)

            assert 'This is a paragraph.' in latex
        finally:
            Path(doc_path).unlink(missing_ok=True)

    def test_multiple_paragraphs(self):
        """测试多个段落"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            doc.add_paragraph('First paragraph')
            doc.add_paragraph('Second paragraph')
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)
            latex = converter.convert_structure(content)

            assert 'First paragraph' in latex
            assert 'Second paragraph' in latex
        finally:
            Path(doc_path).unlink(missing_ok=True)


class TestTableToLatex(unittest.TestCase):
    """测试表格到 LaTeX 的转换"""

    def test_table_to_latex(self):
        """测试表格转为 LaTeX"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = 'A'
            table.cell(0, 1).text = 'B'
            table.cell(1, 0).text = 'C'
            table.cell(1, 1).text = 'D'
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)
            latex = converter.convert_structure(content)

            assert r'\begin{table}' in latex
            assert r'\begin{tabular}' in latex
            assert r'\hline' in latex
            assert 'A' in latex
            assert 'B' in latex
            assert r'\end{tabular}' in latex
            assert r'\end{table}' in latex
        finally:
            Path(doc_path).unlink(missing_ok=True)

    def test_table_with_multiple_rows(self):
        """测试多行表格"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            table = doc.add_table(rows=4, cols=3)
            for r in range(4):
                for c in range(3):
                    table.cell(r, c).text = f'R{r}C{c}'
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)
            latex = converter.convert_structure(content)

            assert r'\begin{table}' in latex
            # 检查行数：应该生成 4 行数据
            assert latex.count(r'\\') >= 4
        finally:
            Path(doc_path).unlink(missing_ok=True)


class TestConvertStructure(unittest.TestCase):
    """测试 convert_structure 方法"""

    def test_mixed_content(self):
        """测试混合内容转换"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        try:
            doc = Document()
            doc.add_heading('Section 1', level=1)
            doc.add_paragraph('Paragraph in section 1')
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = 'Cell 1'
            table.cell(0, 1).text = 'Cell 2'
            table.cell(1, 0).text = 'Cell 3'
            table.cell(1, 1).text = 'Cell 4'
            doc.save(doc_path)

            converter = Docx2LaTeXConverter()
            content = converter.extract_content(doc_path)
            latex = converter.convert_structure(content)

            assert r'\section{Section 1}' in latex
            assert 'Paragraph in section 1' in latex
            assert r'\begin{table}' in latex
        finally:
            Path(doc_path).unlink(missing_ok=True)


class TestSyncConvert(unittest.TestCase):
    """测试同步 convert 方法"""

    def test_convert_basic(self):
        """测试基本转换功能"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc_path = f.name

        with tempfile.TemporaryDirectory() as tmpdir:
            try:
                doc = Document()
                doc.add_heading('Test Document', level=1)
                doc.add_paragraph('Some content here')
                doc.save(doc_path)

                converter = Docx2LaTeXConverter()
                result = converter.convert(doc_path, output_path=Path(tmpdir) / 'output.tex')

                assert result['success'] is True
                assert 'output_path' in result
                assert Path(result['output_path']).exists()
            finally:
                Path(doc_path).unlink(missing_ok=True)


if __name__ == "__main__":
    unittest.main()